"""
将postman导出的文件生成api接口word文档
"""
import os,sys,json
from pprint import pprint 
from docx import Document
from docx.shared import Inches,Pt # 导入字体相关库
from docx.oxml.ns import qn
from docx.shared import RGBColor

postman_file_dict = {}

def init_file():
    """ 获取当前文件夹下所有以postman_collection.json结尾的文件 """
    #当前文件路径
    abp = os.path.split(os.path.realpath(__file__))[0]    
    #过滤当前文件夹下以.postman_collection.json为结尾的文件
    list_file = [f for f in os.listdir(abp) if f.endswith('.postman_collection.json')]
    index = [i for i in range(1,len(list_file)+1)]
    return dict(zip(index,[os.path.join(abp,fp) for fp in list_file]))

def chg_font(obj,fontname='微软雅黑',size=None):
    """设置字体函数"""
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
    if size and isinstance(size,Pt):
        obj.font.size = size


def doc_add_paragraph(doc,text,size=None,color=[],style='Normal'):
    """添加段落并设置字体和颜色"""
    # desc = x['request']["description"] if 'description' in x['request'].keys() else '无'
    # run = doc.add_paragraph().add_run('简要描述：'+desc)
    # run.font.size = Pt(10) # 改变字体大小
    # run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9) #改变字体颜色
    run = doc.add_paragraph(style=style).add_run(text)
    if size!=None:
        run.font.size = Pt(size)
    if color!=[]:
        r,g,b = color
        run.font.color.rgb = RGBColor(r,g,b)

if __name__ == "__main__":
    
    postman_file_dict = init_file()
    print("遍历当前文件夹的postman文件···")
    print("可选文件：")
    for (k,v) in postman_file_dict.items():
        print("%s.%s" % (k,v.split(os.sep)[-1]))

    choice_pos= int(input("%s" % "输入转换为word文档的文件编号："))
    if choice_pos not in postman_file_dict.keys():
        print('编号不在列表内')
        sys.exit(0)
    
    # 反序列化json文件
    with open(postman_file_dict.get(choice_pos),"r",encoding="utf-8") as f:
        item = json.load(f)
    
    doc = Document()
    #设置字体
    chg_font(doc.styles['Normal'],fontname='微软雅黑',size = Pt(10))
    # doc = Document('a.docx')
    doc.add_heading(item['info']['name'],level = 0)

    # 判断数组中是否有该字段
    check_exist = lambda l,k: l[k] if k in l.keys() else ''
    
    # par = doc.add_paragraph('描述：'+item['info']['description'])
    doc_add_paragraph(doc,'文档描述：'+check_exist(item['info'],'description'),13)

    

    for x in item['item']:
        doc.add_paragraph().add_run().add_break()
        p = doc.add_paragraph()
        run = p.add_run('接口：'+x['name'])
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0,205,73)
        
        if 'request' not in x.keys():
            continue
        link = x['request']['url']['raw']
        doc_add_paragraph(doc,'连接：'+link,10,[187, 0, 255],'List Bullet')
        desc = check_exist(x['request'],'description')
        doc_add_paragraph(doc,'简要描述：'+desc,10,[230, 138, 192],'List Bullet')
        doc_add_paragraph(doc,'请求方式：'+x['request']['method'],10,[255,219,0],'List Bullet')

        if 'body' not in x['request'].keys():
            continue
        doc_add_paragraph(doc,'参数：')
        # 写个表格
        table = doc.add_table(1,4)
        hc = table.rows[0].cells
        hc[0].text = '参数名'
        hc[1].text = '参考值'
        hc[2].text = '类型'
        hc[3].text = '说明'
        # 动态生成表格
        for fd in x['request']['body']['formdata']:
            if 'disabled' in fd.keys():
                continue
            cells = table.add_row().cells
            cells[0].text = fd['key']
            cells[1].text = fd['value']
            cells[2].text = fd['type']
            cells[3].text = check_exist(fd,'description')

        # 生成输入示例    
        doc_add_paragraph(doc,'输入示例：',11)
        form = {i['key']:i['value'] for i in x['request']['body']['formdata'] if 'disbale' not in i.keys()}
        exp = json.dumps(form,indent=4,ensure_ascii=False)
        doc_add_paragraph(doc,exp,9.5,[96,125,139])

    abp = os.path.split(os.path.realpath(__file__))[0]
    name =  item['info']['name']+'.docx'
    save_name = os.path.join(abp,name)
    doc.save(save_name)
    print('生成成功！'+ name)
